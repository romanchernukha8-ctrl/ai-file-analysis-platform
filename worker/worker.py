import os
import time
from io import BytesIO
from pathlib import Path

import pika
import psycopg2
import redis
import requests
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas

os.makedirs("/app/generated_files", exist_ok=True)

redis_client = redis.Redis(host="redis", port=6379, decode_responses=True)


def get_db_connection():
    return psycopg2.connect(
        host=os.getenv("POSTGRES_HOST"),
        dbname=os.getenv("POSTGRES_DB"),
        user=os.getenv("POSTGRES_USER"),
        password=os.getenv("POSTGRES_PASSWORD"),
    )


def analyze_with_ollama(text):

    prompt = f"""
Analyze this document.

Return:

1. Short summary
2. Grammar mistakes
3. Suggestions for improvement

Document:

{text[:5000]}
"""

    response = requests.post(
        "http://ollama:11434/api/generate",
        json={"model": "qwen2.5:0.5b", "prompt": prompt, "stream": False},
        timeout=300,
    )
    print("OLLAMA STATUS:", response.status_code, flush=True)
    print("OLLAMA STATUS:", response.text, flush=True)

    # return response.json()["response"]
    data = response.json()

    if "response" not in data:
        raise Exception(str(data))

    return data["response"]


credentials = pika.PlainCredentials("admin", "admin")

connection = None

for attempt in range(10):

    try:

        connection = pika.BlockingConnection(
            pika.ConnectionParameters(host="rabbitmq", credentials=credentials)
        )

        print("Connected to RabbitMQ", flush=True)

        break

    except Exception as e:

        print(f"RabbitMQ not ready ({attempt + 1}/10): {e}", flush=True)

        time.sleep(5)

if connection is None:
    raise Exception("Cannot connect to RabbitMQ")

channel = connection.channel()

channel.queue_declare(queue="file_queue")


def extract_text(file_path):

    ext = Path(file_path).suffix.lower()

    text = ""
    pages = 0

    if ext == ".pdf":

        reader = PdfReader(file_path)

        pages = len(reader.pages)

        for page in reader.pages:
            text += page.extract_text() or ""

    elif ext == ".docx":

        doc = Document(file_path)

        pages = len(doc.paragraphs)

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

    elif ext == ".xlsx":

        workbook = load_workbook(file_path, data_only=True)

        pages = len(workbook.sheetnames)

        for sheet in workbook.worksheets:

            text += f"\n=== SHEET: {sheet.title} ===\n"

            for row in sheet.iter_rows(values_only=True):

                row_text = " ".join(str(cell) for cell in row if cell is not None)

                text += row_text + "\n"

    else:

        raise Exception(f"Unsupported file type: {ext}")

    return text, pages


def add_watermark(input_pdf, output_pdf):

    reader = PdfReader(input_pdf)

    writer = PdfWriter()

    packet = BytesIO()

    can = canvas.Canvas(packet)

    can.setFont("Helvetica", 40)

    can.saveState()

    can.translate(300, 400)

    can.rotate(45)

    can.drawString(0, 0, "AI ANALYZED")

    can.restoreState()

    can.save()

    packet.seek(0)

    watermark = PdfReader(packet)

    watermark_page = watermark.pages[0]

    for page in reader.pages:

        page.merge_page(watermark_page)

        writer.add_page(page)

    os.makedirs(os.path.dirname(output_pdf), exist_ok=True)

    with open(output_pdf, "wb") as f:

        writer.write(f)


def callback(ch, method, properties, body):

    filename = body.decode()

    print(f"Processing: {filename}", flush=True)

    conn = get_db_connection()

    cur = conn.cursor()

    cur.execute(
        """
        SELECT id, filepath
        FROM files
        WHERE filename = %s
        ORDER BY id DESC
        LIMIT 1
        """,
        (filename,),
    )

    row = cur.fetchone()

    if not row:
        print(f"File not found in database: {filename}", flush=True)

        ch.basic_ack(delivery_tag=method.delivery_tag)
        cur.close()
        conn.close()

        return

    file_id = row[0]
    file_path = row[1]

    redis_client.set(f"file:{file_id}:status", "processing")

    file_path = f"/app/{file_path}"

    print(f"FILE PATH: {file_path}", flush=True)

    try:

        text, pages = extract_text(file_path)

    except Exception as e:

        print(f"Processing error: {e}", flush=True)

        cur.execute(
            """
            UPDATE files
            SET status = 'failed'
            WHERE id = %s
            """,
            (file_id,),
        )
        redis_client.set(f"file:{file_id}:status", "failed")

        conn.commit()

        ch.basic_ack(delivery_tag=method.delivery_tag)

        cur.close()
        conn.close()

        return

    text = text.replace("\x00", " ")

    words = len(text.split())

    symbols = len(text)

    print(f"Pages: {pages}, Words: {words}, Symbols: {symbols}", flush=True)
    from pathlib import Path

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        output_pdf = f"/app/generated_files/" f"watermarked_{filename}"

        add_watermark(file_path, output_pdf)

        print(f"Watermark created: {output_pdf}", flush=True)
        watermarked_filename = f"watermarked_{filename}"

        cur.execute(
            """
            UPDATE files
            SET generated_file = %s
            WHERE id = %s
            """,
            (watermarked_filename, file_id),
        )

        conn.commit()

    print("Sending text to Ollama...", flush=True)

    try:

        ai_analysis = analyze_with_ollama(text)

        print("AI RESPONSE:", flush=True)
        print(ai_analysis, flush=True)

    except Exception as e:

        ai_analysis = f"Ollama error: {e}"

        print(ai_analysis, flush=True)

    conn = get_db_connection()

    cur = conn.cursor()

    cur.execute(
        """
        UPDATE files
        SET status = 'processing'
        WHERE filename = %s
        """,
        (filename,),
    )

    conn.commit()

    cur.execute(
        """
        INSERT INTO analysis_results
        (
            file_id,
            pages,
            words,
            symbols
        )
        VALUES (%s, %s, %s, %s)
        """,
        (file_id, pages, words, symbols),
    )

    conn.commit()

    cur.execute(
        """
        INSERT INTO document_texts
        (
            file_id,
            content
        )
        VALUES (%s, %s)
        """,
        (file_id, text),
    )

    conn.commit()

    cur.execute(
        """
        INSERT INTO ai_analysis
        (
            file_id,
            analysis
        )
        VALUES (%s, %s)
        """,
        (file_id, ai_analysis),
    )

    conn.commit()

    cur.execute(
        """
        UPDATE files
        SET status = 'completed'
        WHERE id = %s
        """,
        (file_id,),
    )

    conn.commit()
    redis_client.set(f"file:{file_id}:status", "completed")

    cur.close()
    conn.close()

    ch.basic_ack(delivery_tag=method.delivery_tag)


channel.basic_consume(queue="file_queue", on_message_callback=callback)

print("Waiting for messages...", flush=True)

channel.start_consuming()
